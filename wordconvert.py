#Function to convert the data to a word document
from docx import Document
from docx.shared import RGBColor

def to_word(data, filename):
    doc = Document()

    table = doc.add_table(rows=1, cols=len(data.columns))
    for i, column in enumerate(data.columns):
        table.cell(0, i).text = column

    for _, row in data.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    doc.save(filename)

#Function to convert the result to a word document
def result_word(results_df, output_file):
    doc = Document()
    table = doc.add_table(rows=1, cols=len(results_df.columns))
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(results_df.columns):
        hdr_cells[i].text = column_name

    for _, row in results_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = row_cells[i]
            if isinstance(value, str):
                cell.text = value
            else:
                cell.text = str(value)
            if row['Action'] == 'BUY':
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Green color for BUY
            elif row['Action'] == 'SELL':
                profit_or_loss = row.get('Price', 0) - row.get('Previous Price', 0)
                if profit_or_loss > 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Green color for profit
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for loss

    doc.save(output_file)